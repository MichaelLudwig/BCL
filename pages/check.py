import streamlit as st
import os
from datetime import datetime
from typing import List, Dict, Any
import docx
import json
import re
from openai_api import OpenAIAPI, Dokumenteninfo

st.set_page_config(
    page_title="Check",
    page_icon="📄",
    layout="wide"
)

# OpenAI API Client initialisieren ------------------------------------------------------------------------------------------------------------------------------------------------------------
openai_api = OpenAIAPI()

# Funktion zum Extrahieren der Kapitelstruktur des Brandschutztechnischen Gesamtkonzepts ------------------------------------------------------------------------------------------------------------------------------------------------------------
def extract_chapter_structure(doc: docx.Document) -> Dict[str, Any]:
    """Extrahiert die Kapitelstruktur des Brandschutztechnischen Gesamtkonzepts"""
    
    chapters = []
    current_chapter = None
    current_subchapter = None
    in_brandschutz_chapter = False
    chapter_level = 0
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        # Erkennen der Überschriftenebene
        level = 0
        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name[-1])
            except ValueError:
                continue
                
        text = paragraph.text.strip()
        
        # Hauptkapitel "Brandschutztechnisches Gesamtkonzept" finden
        if level == 1 and "brandschutztechnisches gesamtkonzept" in text.lower():
            in_brandschutz_chapter = True
            current_chapter = {"title": text, "content": "", "subchapters": []}
            chapters.append(current_chapter)
            chapter_level = 1
            continue
            
        # Wenn wir das nächste Hauptkapitel erreichen, beenden wir die Extraktion
        if in_brandschutz_chapter and level == 1 and current_chapter and "brandschutztechnisches gesamtkonzept" not in text.lower():
            break
            
        if in_brandschutz_chapter:
            if level == 2:  # Unterkapitel
                # Extrahiere die Kapitelnummer, falls vorhanden
                chapter_number = ""
                title = text
                # Suche nach Mustern wie "5.1" oder "5" am Anfang des Textes
                number_match = re.match(r'^(\d+(\.\d+)?)\s+(.+)$', text)
                if number_match:
                    chapter_number = number_match.group(1)
                    title = number_match.group(3)
                
                current_subchapter = {
                    "number": chapter_number,
                    "title": title,
                    "content": "",
                    "subchapters": []
                }
                current_chapter["subchapters"].append(current_subchapter)
                chapter_level = 2
            elif level == 3:  # Unter-Unterkapitel
                if current_subchapter:
                    # Extrahiere die Kapitelnummer, falls vorhanden
                    chapter_number = ""
                    title = text
                    # Suche nach Mustern wie "5.1.1" am Anfang des Textes
                    number_match = re.match(r'^(\d+(\.\d+){1,2})\s+(.+)$', text)
                    if number_match:
                        chapter_number = number_match.group(1)
                        title = number_match.group(3)
                    
                    current_subchapter["subchapters"].append({
                        "number": chapter_number,
                        "title": title,
                        "content": "",
                        "subchapters": []
                    })
                    chapter_level = 3
            elif level == 0:  # Inhalt
                if chapter_level == 1 and current_chapter:
                    current_chapter["content"] += text + "\n"
                elif chapter_level == 2 and current_subchapter:
                    current_subchapter["content"] += text + "\n"
                elif chapter_level == 3 and current_subchapter and current_subchapter["subchapters"]:
                    current_subchapter["subchapters"][-1]["content"] += text + "\n"
    
    return {"brandschutzkonzept": chapters[0] if chapters else None}

# Funktion zum Extrahieren des Dokumenteninhalts ------------------------------------------------------------------------------------------------------------------------------------------------------------
def extract_document_content(file):
    if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Temporäre Datei speichern
        with open("temp.docx", "wb") as f:
            f.write(file.getvalue())
        
        # Word-Dokument einlesen
        doc = docx.Document("temp.docx")
        
        # Extrahiere die Kapitelstruktur und speichere sie im Session State
        chapter_structure = extract_chapter_structure(doc)
        st.session_state['chapter_structure'] = chapter_structure
        
        # Text aus dem Dokument extrahieren
        content = ""
        
        # Extrahiere Text aus Tabellen der ersten Seite
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text.strip() + " | "
                content += "\n"
            content += "\n"
        
        # Extrahiere Text aus Absätzen
        for i, paragraph in enumerate(doc.paragraphs):
            if i < 200:  # Erhöhe auf 150 Absätze
                # Überspringe leere Absätze
                if paragraph.text.strip():
                    # Füge Formatierungsinformationen hinzu
                    if paragraph.style.name.startswith('Heading'):
                        content += f"\n### {paragraph.text}\n"
                    else:
                        content += paragraph.text + "\n"
        
        # Temporäre Datei löschen
        os.remove("temp.docx")
                
        return content
    else:
        st.error("Bitte laden Sie ein Word-Dokument (.docx) hoch")
        return None

# Seitenleisten-Konfiguration ------------------------------------------------------------------------------------------------------------------------------------------------------------
with st.sidebar:
    st.title("Dokument Upload")
    
    # Authentication Switch
    st.write("### Authentifizierung")
    use_managed_identity = st.toggle("Managed Identity verwenden", value=True, help="Deaktivieren für lokale Tests mit API Key")
    if not use_managed_identity:
        st.info("Lokaler Modus: Verwendet API Key Authentifizierung")
    
    # Environment Variable setzen
    os.environ['WEBSITE_INSTANCE_ID'] = '1' if use_managed_identity else ''
    
    # Suchfilter ------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.write("### Suchfilter")
    
    # Bundesländer Multiselect
    bundeslaender = [
        "Baden-Württemberg", "Bayern", "Berlin", "Brandenburg", "Bremen",
        "Hamburg", "Hessen", "Mecklenburg-Vorpommern", "Niedersachsen",
        "Nordrhein-Westfalen", "Rheinland-Pfalz", "Saarland", "Sachsen",
        "Sachsen-Anhalt", "Schleswig-Holstein", "Thüringen"
    ]
    selected_bundeslaender = st.multiselect(
        "Bundesländer",
        options=bundeslaender,
        default=[]
    )
    
    # Baukategorien Multiselect
    baukategorien = [
        "Bauordnung", "Beherbergungsstätten", "elektrische Betriebsräume",
        "FeuerungsVO", "Flächen für Feuerwehr", "Garagen", "Hochhaus",
        "Industriebau", "Krankenhaus", "Leitungsanlagen", "Lüftungsanlagen",
        "Schulen", "Systemboden", "Verkaufsstätte", "Versammlungsstätte",
        "VwV TechnischeBaubetimmungen"
    ]
    selected_baukategorien = st.multiselect(
        "Baukategorien",
        options=baukategorien,
        default=[]
    )
    
    # Suchfilter anzeigen
    st.write("### Generierter Suchfilter")
    filter_conditions = []
    search_filter = None  # Initialisiere search_filter als None
    
    if selected_bundeslaender:
        bundeslaender_filter = " or ".join([f"Bundesland eq '{land}'" for land in selected_bundeslaender])
        filter_conditions.append(f"({bundeslaender_filter})")
    
    if selected_baukategorien:
        baukategorien_filter = " or ".join([f"Baukategorie eq '{kat}'" for kat in selected_baukategorien])
        filter_conditions.append(f"({baukategorien_filter})")
    
    if filter_conditions:
        search_filter = " and ".join(filter_conditions)
        st.text_area("Aktueller Filter", value=search_filter, height=100, disabled=True)
    else:
        st.text_area("Aktueller Filter", value="Kein Filter aktiv", height=100, disabled=True)
    
    # Speichere den Filter im Session State für spätere Verwendung
    st.session_state['search_filter'] = search_filter
    
    # File Uploader in der Seitenleiste ------------------------------------------------------------------------------------------------------------------------------------------------------------
    st.write("### Dokument hochladen")
    uploaded_file = st.file_uploader("Word-Dokument hochladen", type=['docx'])
    
    if uploaded_file is not None:
        # Anzeige des Dateinamens
        st.success(f"Hochgeladene Datei: {uploaded_file.name}")
        
        # Zusätzliche Datei-Informationen
        file_details = {
            "Dateiname": uploaded_file.name,
            "Dateigröße": f"{uploaded_file.size / 1024:.2f} KB",
            "Zeitpunkt": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        

# Hauptbereich ------------------------------------------------------------------------------------------------------------------------------------------------------------
st.title("Dokument Überprüfung")

if uploaded_file is None:
    st.info("Bitte laden Sie ein Word-Dokument (.docx) in der Seitenleiste hoch.")
else:
    try:
        # Prüfe ob das Dokument bereits analysiert wurde
        if ('current_file' not in st.session_state or 
            st.session_state['current_file'] != uploaded_file.name or 
            'doc_info' not in st.session_state):
            
            # Dokument analysieren
            with st.spinner("Dokument wird analysiert..."):
                content = extract_document_content(uploaded_file)
                if content:
                    # OpenAI API aufrufen mit Filter
                    doc_info = openai_api.extract_document_info(
                        content, 
                        search_filter=st.session_state.get('search_filter')
                    )
                    # Speichere die Ergebnisse im Session State
                    st.session_state['current_file'] = uploaded_file.name
                    st.session_state['doc_info'] = doc_info
        else:
            # Verwende die gespeicherten Informationen
            doc_info = st.session_state['doc_info']
        
        # Wenn doc_info verfügbar ist (entweder neu analysiert oder aus dem Cache)
        if 'doc_info' in st.session_state:
            doc_info = st.session_state['doc_info']
            # Anzeige der extrahierten Informationen
            with st.expander("Dokumentinformationen", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**Bauvorhaben:**")
                    st.write(doc_info.bauvorhaben)
                    st.write("**Bauherr:**")
                    st.write(doc_info.bauherr)
                    st.write("**Entwurfsverfasser:**")
                    st.write(doc_info.entwurfsverfasser) 
                
                with col2:
                    st.write("**Bundesland:**")
                    st.write(doc_info.bundesland)
                    st.write("**Erstellt am:**")
                    st.write(doc_info.erstellt_am)
                    st.write("**Zielstellung:**")
                    st.write(doc_info.zielstellung)
                
                st.write("**Bauordnungsrechtliche Grundlagen:**")
                for grundlage in doc_info.bauordnungsrechtliche_grundlagen:
                    st.write(f"- {grundlage}")
            
            # Kapitelstruktur anzeigen 
            st.write("# Brandschutztechnisches Gesamtkonzept")
            
            if 'chapter_structure' in st.session_state and st.session_state['chapter_structure']['brandschutzkonzept']:
                chapter = st.session_state['chapter_structure']['brandschutzkonzept']
                
                # Container für besseres Styling
                with st.container():
                    # Unterkapitel
                    for i, subchapter in enumerate(chapter['subchapters']):
                        # Container für Unterkapitel mit Prüf-Button und Inhalt
                        with st.container():
                            # Überschrift und Prüf-Button 
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                chapter_title = f"{subchapter['number']} {subchapter['title']}" if subchapter.get('number') else subchapter['title']
                                st.markdown(f"## {chapter_title}")
                            with col2:
                                if st.button("Prüfen", key=f"check_{i}"):
                                    if 'chapters_data' not in st.session_state:
                                        st.session_state['chapters_data'] = {}
                                    
                                    chapter_key = subchapter['title']
                                    if chapter_key not in st.session_state['chapters_data']:
                                        st.session_state['chapters_data'][chapter_key] = {
                                            'title': subchapter['title'],
                                            'content': subchapter['content'],
                                            'subchapters': subchapter['subchapters'],
                                            'report': '',
                                            'is_selected': False
                                        }
                                    
                                    # Setze alle Kapitel auf nicht ausgewählt
                                    for key in st.session_state['chapters_data']:
                                        st.session_state['chapters_data'][key]['is_selected'] = (key == chapter_key)
                                    
                                    # Führe die KI-Prüfung durch
                                    try:
                                        with st.spinner(f"Prüfe Kapitel '{chapter_key}'..."):
                                            # Sammle alle Unterkapitel mit Inhalt
                                            chapters_to_check = []
                                            chapter_mapping = {}  # Neue Map für die Zuordnung
                                            
                                            # Füge Hauptkapitelinhalt hinzu, falls vorhanden
                                            if subchapter['content'].strip():
                                                chapters_to_check.append({
                                                    'title': subchapter['title'],
                                                    'number': subchapter.get('number', ''),
                                                    'content': subchapter['content']
                                                })
                                                chapter_mapping[len(chapters_to_check) - 1] = f"{i}_main"
                                            
                                            # Füge Unterkapitel hinzu
                                            for j, subsubchapter in enumerate(subchapter['subchapters']):
                                                if subsubchapter['content'].strip():
                                                    chapters_to_check.append({
                                                        'title': subsubchapter['title'],
                                                        'number': subsubchapter.get('number', ''),
                                                        'content': subsubchapter['content']
                                                    })
                                                    chapter_mapping[len(chapters_to_check) - 1] = f"{i}_{j}"
                                            
                                            if chapters_to_check:
                                                # Prüfe alle Unterkapitel in einer Anfrage
                                                result = openai_api.check_chapter(
                                                    chapters_to_check,
                                                    doc_info
                                                )
                                                
                                                # Wenn ein Fehler aufgetreten ist
                                                if "error" in result:
                                                    st.error(result["error"])
                                                    continue
                                                
                                                # Initialisiere subchapter_reports im Session State
                                                if 'subchapter_reports' not in st.session_state['chapters_data'][chapter_key]:
                                                    st.session_state['chapters_data'][chapter_key]['subchapter_reports'] = {}
                                                
                                                # Speichere die Prüfberichte für jedes Unterkapitel
                                                for j, (chapter, report) in enumerate(zip(chapters_to_check, result['reports'])):
                                                    # Verwende die Mapping-Tabelle für die korrekte Zuordnung
                                                    subchapter_key = chapter_mapping[j]
                                                    
                                                    st.session_state['chapters_data'][chapter_key]['subchapter_reports'][subchapter_key] = {
                                                        'report': report,
                                                        'citations': result['citations']
                                                    }
                                                
                                                # Zeige Token-Nutzung an
                                                st.info(f"Token-Nutzung - Prompt: {result.get('usage', {}).get('prompt_tokens', 'N/A')} | Response: {result.get('usage', {}).get('completion_tokens', 'N/A')}")
                                            
                                    except Exception as e:
                                        st.error(f"Fehler bei der Prüfung: {str(e)}")
                                
                            # Zeige Hauptkapitelinhalt, falls vorhanden 
                            if subchapter['content'].strip():
                                st.markdown("### Hauptkapitelinhalt")
                                main_content_col, main_report_col = st.columns([1, 1])
                                
                                with main_content_col:
                                    st.text_area(
                                        "Kapitelinhalt",
                                        value=subchapter['content'],
                                        height=300,
                                        key=f"maincontent_{i}"
                                    )
                                
                                with main_report_col:
                                    # Hole den Prüfbericht für das Hauptkapitel
                                    mainreport_value = ""
                                    if 'chapters_data' in st.session_state and \
                                       subchapter['title'] in st.session_state['chapters_data'] and \
                                       'subchapter_reports' in st.session_state['chapters_data'][subchapter['title']]:
                                        main_key = f"{i}_main"
                                        if main_key in st.session_state['chapters_data'][subchapter['title']]['subchapter_reports']:
                                            mainreport_value = st.session_state['chapters_data'][subchapter['title']]['subchapter_reports'][main_key]['report']
                                    
                                    st.text_area(
                                        "Prüfbericht",
                                        value=mainreport_value,
                                        height=300,
                                        key=f"mainreport_{i}"
                                    )

                            # Zeige Unterkapitel
                            for j, subsubchapter in enumerate(subchapter['subchapters']):
                                if subsubchapter['content'].strip():
                                    subchapter_title = f"{subsubchapter['number']} {subsubchapter['title']}" if subsubchapter['number'] else subsubchapter['title']
                                    st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;**{subchapter_title}**")
                                    
                                    sub_content_col, sub_report_col = st.columns([1, 1])
                                    
                                    with sub_content_col:
                                        st.text_area(
                                            "Unterkapitelinhalt",
                                            value=subsubchapter['content'],
                                            height=300,
                                            key=f"subcontent_{i}_{j}"
                                        )
                                    
                                    with sub_report_col:
                                        # Hole den Prüfbericht für das Unterkapitel
                                        subreport_value = ""
                                        if 'chapters_data' in st.session_state and \
                                           subchapter['title'] in st.session_state['chapters_data'] and \
                                           'subchapter_reports' in st.session_state['chapters_data'][subchapter['title']]:
                                            subchapter_key = f"{i}_{j}"
                                            if subchapter_key in st.session_state['chapters_data'][subchapter['title']]['subchapter_reports']:
                                                subreport_value = st.session_state['chapters_data'][subchapter['title']]['subchapter_reports'][subchapter_key]['report']
                                        
                                        st.text_area(
                                            "Prüfbericht",
                                            value=subreport_value,
                                            height=300,
                                            key=f"subreport_{i}_{j}"
                                        )
                                
                                # st.markdown("---")  # Trennlinie zwischen den Kapiteln
                            
                            # Nach der Schleife für die Unterkapitel
                            # Zeige die verwendeten Quellen am Ende des Hauptkapitels
                            if 'chapters_data' in st.session_state and \
                               subchapter['title'] in st.session_state['chapters_data'] and \
                               'subchapter_reports' in st.session_state['chapters_data'][subchapter['title']] and \
                               any(report.get('citations') for report in st.session_state['chapters_data'][subchapter['title']]['subchapter_reports'].values()):
                                with st.expander("Verwendete Quellen", expanded=False):
                                    # Sammle alle eindeutigen Zitate
                                    all_citations = []
                                    for report in st.session_state['chapters_data'][subchapter['title']]['subchapter_reports'].values():
                                        if report.get('citations'):
                                            all_citations.extend(report['citations'])
                                    # Entferne Duplikate und zeige die Quellen
                                    unique_citations = list({str(citation): citation for citation in all_citations}.values())
                                    st.json(unique_citations)
                            
                            st.markdown("---")  # Abschließende Trennlinie für das Hauptkapitel
                    
    except Exception as e:
        st.error(f"Fehler bei der Analyse: {str(e)}")
        st.error("Details zum Fehler für Entwickler:")
        st.code(str(e))
 